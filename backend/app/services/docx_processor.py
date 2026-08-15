"""
DOCX Document Processor.

Handles extraction of text from DOCX files and production of
redacted DOCX documents while preserving formatting.
"""

from docx import Document
from docx.shared import RGBColor
import docx
import copy
import io
import re
from typing import Optional

from ..models import PIIEntity


def extract_text(file_bytes: bytes) -> str:
    """
    Extract all text from a DOCX file.

    Concatenates text from paragraphs and table cells,
    preserving paragraph breaks.

    Args:
        file_bytes: Raw bytes of the DOCX file

    Returns:
        The full extracted text
    """
    doc = Document(io.BytesIO(file_bytes))
    parts: list[str] = []

    # Extract from paragraphs
    for paragraph in doc.paragraphs:
        parts.append(paragraph.text)

    # Extract from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    parts.append(text)

    return '\n'.join(parts)


def create_redacted_docx(
    original_bytes: bytes,
    entities: list[PIIEntity],
) -> bytes:
    """
    Create a redacted copy of a DOCX file.

    Replaces PII in the document text with the synthetic replacements
    from the entity list, preserving the original formatting as much
    as possible.

    Args:
        original_bytes: Raw bytes of the original DOCX file
        entities: List of PIIEntity with original and replacement values

    Returns:
        Bytes of the redacted DOCX file
    """
    doc = Document(io.BytesIO(original_bytes))

    # Build a replacement map: original text -> replacement text
    replacement_map: dict[str, str] = {}
    for entity in entities:
        replacement_map[entity.original] = entity.replacement

    # Sort replacements by length (longest first) to avoid partial replacements
    sorted_replacements = sorted(
        replacement_map.items(),
        key=lambda x: len(x[0]),
        reverse=True,
    )

    # Process paragraphs
    for paragraph in doc.paragraphs:
        _replace_in_paragraph(paragraph, sorted_replacements)

    # Process table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_in_paragraph(paragraph, sorted_replacements)

    # Process headers and footers
    for section in doc.sections:
        if section.header:
            for paragraph in section.header.paragraphs:
                _replace_in_paragraph(paragraph, sorted_replacements)
        if section.footer:
            for paragraph in section.footer.paragraphs:
                _replace_in_paragraph(paragraph, sorted_replacements)

    # Redact hyperlink target URLs (href) across all parts
    def _redact_rels_in_part(part):
        for rel in part.rels.values():
            if rel.reltype == docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK:
                target = rel.target_ref
                if target:
                    new_target = target
                    for original, replacement in sorted_replacements:
                        new_target = new_target.replace(original, replacement)
                    if new_target != target:
                        rel._target = new_target

    _redact_rels_in_part(doc.part)
    for section in doc.sections:
        if section.header:
            _redact_rels_in_part(section.header.part)
        if section.footer:
            _redact_rels_in_part(section.footer.part)

    # Save to bytes
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output.read()


def _replace_in_paragraph(
    paragraph,
    replacements: list[tuple[str, str]],
):
    """
    Replace PII text within a paragraph while trying to preserve
    run-level formatting.

    This handles the tricky case where a PII value might be split
    across multiple runs in the DOCX XML. We concatenate runs,
    perform replacement, and then rewrite the runs.
    """
    # Get the full paragraph text
    full_text = paragraph.text
    if not full_text:
        return

    # Check if any replacements apply to this paragraph
    needs_replacement = False
    for original, replacement in replacements:
        if original in full_text:
            needs_replacement = True
            break

    if not needs_replacement:
        return

    # Build a map of character positions to runs
    # Use XPath to find all w:r elements, including those nested in w:hyperlink
    from docx.text.run import Run
    runs_xml = paragraph._p.xpath('.//w:r')
    runs = [Run(r, paragraph) for r in runs_xml]
    if not runs:
        return

    # Concatenate all run texts and track boundaries
    run_texts: list[str] = []
    run_starts: list[int] = []
    pos = 0
    for run in runs:
        run_starts.append(pos)
        run_texts.append(run.text or "")
        pos += len(run.text or "")

    # Perform replacements on the full text
    new_text = full_text
    for original, replacement in replacements:
        new_text = new_text.replace(original, replacement)

    # If nothing changed, skip
    if new_text == full_text:
        return

    # Simple strategy: keep first run's formatting, put all text in first run,
    # clear remaining runs
    if runs:
        runs[0].text = new_text
        for run in runs[1:]:
            run.text = ""


def create_comparison_text(
    original_text: str,
    entities: list[PIIEntity],
) -> str:
    """
    Create a redacted version of the text where PII is replaced
    with synthetic values.

    Args:
        original_text: The original document text
        entities: List of detected PII entities with replacements

    Returns:
        The redacted text
    """
    # Sort entities by position (reverse) to replace from end to start
    # This preserves character offsets
    sorted_entities = sorted(entities, key=lambda e: e.start, reverse=True)

    redacted = original_text
    for entity in sorted_entities:
        redacted = (
            redacted[:entity.start]
            + entity.replacement
            + redacted[entity.end:]
        )

    return redacted
